"""
Fase 5 (`factory/docs/document_remediation_evolution/DOCUMENT_REMEDIATION_SPEC.md`
§2-6, `IMPLEMENTATION_ROADMAP.md`) — documento candidato completo + redline
real, generados a partir de la representación intermedia de Fase 4
(`regulatory/document_structure_extractor.py`) y de N `RemediationChange`
ya validados (mismo dict shape que valida `remediation_package_schemas.py`).

Invariante no negociable (§1): el original NUNCA se toca -- este módulo
solo LEE el PDF fuente (vía la representación ya extraída) y ESCRIBE
documentos `.docx` nuevos en rutas de salida separadas.

Dos artefactos, mismo patrón ya usado en los 2 paquetes reales de la
sesión anterior (`candidate_document` + `redline_document` como
`ArtifactReference` separados):
  - `generate_candidate_document`: version limpia, texto insertado sin
    marcado visual -- el documento final "como quedaria".
  - `generate_redline_document`: version marcada (verde + tag
    `[change_id]`, mismo patron `RGBColor` de `gmpai_docx_draft.py`) --
    la que ancla el chequeo `DOCUMENT_CONFORMANCE` automatizado (§5),
    porque es la unica con marcas explicitas y verificables por posicion.

Solo implementa `CONTENT_ADDITION` -- los 3 cambios reales existentes
(COR-1, COR-2, COR-5, ambos paquetes de la sesión anterior) son los
únicos casos reales, y los 3 son `CONTENT_ADDITION`. `CONTENT_REPLACEMENT`
existe en el schema pero no tiene caso real que lo ejercite todavía --
se declara `NotImplementedError` explícito (fail-closed) en vez de
adivinar un comportamiento de reemplazo sin evidencia real.
"""
from __future__ import annotations

import hashlib

from docx import Document
from docx.shared import Pt, RGBColor

_INSERTED_COLOR = RGBColor(0x00, 0x70, 0x00)
_TAG_COLOR = RGBColor(0x60, 0x60, 0x60)
_WARNING_COLOR = RGBColor(0xB0, 0x40, 0x00)


def _sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _validate_change_type(change: dict) -> None:
    if change["change_type"] != "CONTENT_ADDITION":
        raise NotImplementedError(
            f"change_type={change['change_type']!r} (change_id={change['change_id']!r}) "
            "-- solo CONTENT_ADDITION tiene caso real validado hoy; "
            "CONTENT_REPLACEMENT no se implementa sin evidencia real que lo ejercite."
        )


def _section_for_page(structure: dict, page_start: int) -> dict:
    """La última sección de nivel 1 cuya pagina_inicio <= page_start
    (misma sección donde cae esa página real del documento)."""
    candidate = None
    for seccion in structure["secciones"]:
        if seccion["pagina_inicio"] <= page_start:
            candidate = seccion
        else:
            break
    if candidate is None:
        raise ValueError(
            f"page_start={page_start} es anterior a la primera sección detectada "
            "-- no hay dónde anclar el cambio."
        )
    return candidate


def _group_changes_by_section(structure: dict, changes: list[dict]) -> dict[str, list[dict]]:
    grouped: dict[str, list[dict]] = {}
    for change in changes:
        _validate_change_type(change)
        page_start = change["citations"][0]["page_start"]
        seccion = _section_for_page(structure, page_start)
        grouped.setdefault(seccion["numero"], []).append(change)
    return grouped


def _add_warning(doc: Document, texto: str) -> None:
    warn = doc.add_paragraph()
    run = warn.add_run(texto)
    run.bold = True
    run.font.color.rgb = _WARNING_COLOR
    run.font.size = Pt(10)


def generate_candidate_document(structure: dict, changes: list[dict]) -> Document:
    """Documento candidato completo, version limpia (§2): texto original
    intacto + `proposed_content` de cada cambio insertado en su sección,
    sin marcado visual de diff -- el documento "como quedaria"."""
    doc = Document()
    doc.add_heading("Documento candidato completo -- DRAFT", level=1)
    _add_warning(
        doc,
        "ESTADO: DRAFT. Los cambios documentales NO garantizan por si solos "
        "cumplimiento final -- requiere revision QA/Validacion humana "
        "(IMPLEMENTATION_VERIFICATION y REGULATORY_COMPLIANCE quedan SIEMPRE "
        "fuera de este sistema). El original nunca fue modificado "
        "(SOURCE_IMMUTABLE).",
    )

    for texto in structure.get("texto_previo_a_primera_seccion", []):
        doc.add_paragraph(texto)

    grouped = _group_changes_by_section(structure, changes)
    for seccion in structure["secciones"]:
        doc.add_heading(f"{seccion['numero']} {seccion['titulo']}", level=2)
        for texto in seccion["parrafos"]:
            doc.add_paragraph(texto)
        for change in grouped.get(seccion["numero"], []):
            doc.add_paragraph(change["proposed_content"])

    return doc


def generate_redline_document(structure: dict, changes: list[dict]) -> tuple[Document, list[dict]]:
    """Redline real (§3): mismo documento, pero el texto insertado queda
    marcado (verde + tag `[change_id]`, mismo patron RGBColor de
    `gmpai_docx_draft.py`). Devuelve (Document, insertion_manifest) --
    insertion_manifest: [{change_id, section_numero, paragraph_index,
    proposed_content_sha256}], donde `paragraph_index` es la posición
    real dentro de `Document.paragraphs` (incluye título/warning/headings,
    contado con `len(doc.paragraphs) - 1` tras cada inserción para no
    depender de un conteo manual paralelo que pueda desincronizarse)."""
    doc = Document()
    doc.add_heading("Documento candidato -- REDLINE (DRAFT)", level=1)
    _add_warning(
        doc,
        "ESTADO: DRAFT. Texto en verde con tag [change_id] = insertado por "
        "remediacion documental, NO presente en el original. Revision "
        "QA/Validacion humana requerida antes de cualquier uso.",
    )

    for texto in structure.get("texto_previo_a_primera_seccion", []):
        doc.add_paragraph(texto)

    grouped = _group_changes_by_section(structure, changes)
    insertion_manifest: list[dict] = []
    for seccion in structure["secciones"]:
        doc.add_heading(f"{seccion['numero']} {seccion['titulo']}", level=2)
        for texto in seccion["parrafos"]:
            doc.add_paragraph(texto)
        for change in grouped.get(seccion["numero"], []):
            p = doc.add_paragraph()
            tag_run = p.add_run(f"[{change['change_id']}] ")
            tag_run.bold = True
            tag_run.italic = True
            tag_run.font.color.rgb = _TAG_COLOR
            content_run = p.add_run(change["proposed_content"])
            content_run.font.color.rgb = _INSERTED_COLOR
            content_run.font.underline = True
            insertion_manifest.append({
                "change_id": change["change_id"],
                "section_numero": seccion["numero"],
                "paragraph_index": len(doc.paragraphs) - 1,
                "proposed_content_sha256": _sha256_text(change["proposed_content"]),
            })

    return doc, insertion_manifest


def verify_document_conformance(
    docx_path: str, changes: list[dict], insertion_manifest: list[dict]
) -> list[dict]:
    """`DOCUMENT_CONFORMANCE` (SPEC §5, la única de las 3 revalidaciones
    automatizable por este sistema): reabre el `.docx` YA GUARDADO en
    disco (nunca el objeto `Document` en memoria -- un chequeo contra
    memoria no detectaría un error real de serialización) y verifica que
    el texto insertado en `paragraph_index` (declarado por
    `insertion_manifest`, producido por `generate_redline_document`)
    coincide byte a byte con `proposed_content`.

    `IMPLEMENTATION_VERIFICATION` y `REGULATORY_COMPLIANCE` (SPEC §5)
    permanecen SIEMPRE fuera de este chequeo -- nunca automatizables por
    este sistema, ninguna función de este módulo las toca."""
    reopened = Document(docx_path)
    paragraphs = reopened.paragraphs
    manifest_by_change_id = {m["change_id"]: m for m in insertion_manifest}

    results: list[dict] = []
    for change in changes:
        change_id = change["change_id"]
        entry = manifest_by_change_id.get(change_id)
        if entry is None:
            results.append({
                "change_id": change_id, "status": "CHANGE_NOT_APPLIED",
                "reason": "sin entrada en insertion_manifest",
            })
            continue

        idx = entry["paragraph_index"]
        if idx >= len(paragraphs):
            results.append({
                "change_id": change_id, "status": "CHANGE_NOT_APPLIED",
                "reason": f"paragraph_index {idx} fuera de rango ({len(paragraphs)} parrafos en el docx reabierto)",
            })
            continue

        inserted_runs = [
            run.text for run in paragraphs[idx].runs
            if run.font.color and run.font.color.rgb == _INSERTED_COLOR
        ]
        actual_text = "".join(inserted_runs)
        expected_text = change["proposed_content"]
        if actual_text == expected_text and _sha256_text(actual_text) == entry["proposed_content_sha256"]:
            results.append({
                "change_id": change_id, "status": "DOCUMENT_CONFORMANCE",
                "paragraph_index": idx,
            })
        else:
            results.append({
                "change_id": change_id, "status": "CHANGE_NOT_APPLIED",
                "reason": "texto insertado no coincide byte a byte con proposed_content",
                "paragraph_index": idx,
            })

    return results
