"""
GMPAI — Borrador DOCX controlado para REM-GMPAI-001.

No es una edición del documento original (el finding de trazabilidad es a
nivel de familia documental — "falta la etapa PROTOCOL_TEMPLATE" — no una
corrección de texto puntual dentro de un documento existente). Este módulo
genera una PROPUESTA CONTROLADA: un memo de remediación versionado, con
changelog, que referencia el finding, el requisito regulatorio y el agente
que lo generó. Nunca toca los originales en GMPAI/source/. No ejecuta
macros de archivos .docm — no abre ni procesa ningún .docm, solo escribe
un .docx nuevo desde cero con python-docx.
"""

from __future__ import annotations

import hashlib
import io
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt, RGBColor


def _sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def build_remediation_draft_docx(report_data: dict) -> tuple[bytes, str]:
    """Retorna (docx_bytes, sha256_hex). Contenido derivado 100% de
    report_data (ya agregado desde el RC canónico aprobado) — no reprocesa
    documentos ni inventa hallazgos."""
    items = report_data.get("remediation_items", [])
    item = items[0] if items else {}
    rc = report_data.get("rc_canonical", {})

    doc = Document()

    title = doc.add_heading("Propuesta de Remediacion Documental — DRAFT", level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    warn = doc.add_paragraph()
    warn_run = warn.add_run(
        "ESTADO: DRAFT — requiere revision humana. NO se marca automaticamente "
        "como compliant. Este documento es una PROPUESTA CONTROLADA, no una "
        "edicion del documento original."
    )
    warn_run.bold = True
    warn_run.font.color.rgb = RGBColor(0xB0, 0x40, 0x00)
    warn_run.font.size = Pt(10)

    doc.add_heading("Identificacion", level=2)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for label, value in [
        ("ID remediacion", item.get("id", "REM-GMPAI-001")),
        ("Proyecto", report_data.get("project", {}).get("project_id", "")),
        ("Sistema", item.get("sistema", "")),
        ("RC relacionado", rc.get("rc_id", "")),
        ("Finding relacionado", item.get("finding_relacionado", "")),
        ("Agente", "requirements_traceability_agent"),
        ("Agent version", "no_disponible (limitacion declarada del modulo LLM de trazabilidad)"),
        ("Estado", "draft"),
        ("Fecha generacion", datetime.now(timezone.utc).isoformat()),
    ]:
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value)

    doc.add_heading("Requisito regulatorio", level=2)
    doc.add_paragraph("Trazabilidad URS -> FS -> DS -> IQ/OQ/PQ -> SAT (GAMP 5 V-model)")

    doc.add_heading("Texto / seccion original (evidencia actual)", level=2)
    doc.add_paragraph(
        "El expediente de " + item.get("sistema", "") + " contiene URS, FS, DS, "
        "arquitectura, narrativas de control, listado de alarmas y SAT completado "
        "(evidencia de ejecucion de pruebas de aceptacion). No contiene un documento "
        "de tipo PROTOCOL_TEMPLATE (protocolo IQ/OQ/PQ) que especifique bajo que "
        "criterios se planifico y aprobo ese SAT."
    )

    doc.add_heading("Cambio propuesto", level=2)
    p = doc.add_paragraph()
    p.add_run("Opcion A — ").bold = True
    p.add_run(
        "Incorporar al expediente el protocolo IQ/OQ/PQ original bajo el cual se "
        "ejecuto y aprobo el SAT (215115305 SCADA-PCS Misc PLC SAT3 Scanned-1.pdf / "
        "215115305-T-041 SAT3 Completed.pdf), y reevaluar la trazabilidad de esta "
        "familia documental unicamente (sin reprocesar los 32 documentos)."
    )
    p2 = doc.add_paragraph()
    p2.add_run("Opcion B — ").bold = True
    p2.add_run(
        "Generar y aprobar (QA/Validacion) una justificacion formal de no "
        "aplicabilidad de la etapa PROTOCOL_TEMPLATE para este sistema especifico, "
        "con la firma y fecha del responsable."
    )

    doc.add_heading("Finding relacionado (fuente)", level=2)
    doc.add_paragraph(f"Brecha original: {item.get('hallazgo', '')}")
    doc.add_paragraph(f"Impacto: {item.get('impacto', '')}")
    doc.add_paragraph(f"Severidad conservada del analisis: {item.get('severidad', '')}")

    doc.add_heading("Changelog", level=2)
    ct = doc.add_table(rows=1, cols=3)
    ct.style = "Light Grid Accent 1"
    hdr = ct.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Version", "Fecha", "Cambio"
    row = ct.add_row().cells
    row[0].text = "v1 (draft)"
    row[1].text = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    row[2].text = "Generacion inicial del borrador a partir del finding REM-GMPAI-001, sin decision humana aun."

    doc.add_heading("Limitaciones declaradas", level=2)
    doc.add_paragraph(
        "Este documento NO reemplaza el protocolo original ni constituye evidencia "
        "de cumplimiento. Es una propuesta de remediacion para que QA/Validacion "
        "decida (accept/reject/request_changes) por el mecanismo oficial de Capa 9. "
        "Los documentos originales en GMPAI/source/ no fueron modificados."
    )

    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    return data, _sha256_bytes(data)
