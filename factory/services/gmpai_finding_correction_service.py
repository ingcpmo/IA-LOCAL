"""
GMPAI — Clasificacion finding -> correccion y matriz completa (REM-GMPAI-001+
en sentido amplio: no solo el gap de trazabilidad, sino los 267 findings del
RC canonico).

Este modulo es una capa de CLASIFICACION DETERMINISTA sobre findings YA
existentes en el RC canonico aprobado (pipeline_pilot_llm.json vía
gmpai_artifact_service.build_final_report_data()). NO invoca agentes, NO
llama a Ollama, NO reprocesa documentos: solo aplica una regla fija por
`estado` para decidir si un finding es corregible, requiere evidencia, o
requiere decision humana — y arma finding_id + matriz para el reporte.

Reglas (ver auditoria REM-GMPAI-001 / seccion 4 del encargo):
  - cumple_parcialmente -> correction_generated (completa/mejora la seccion,
    conserva el contenido valido; el texto propuesto es la recomendacion ya
    generada por el agente, nunca un hecho nuevo inventado).
  - no_cumple -> correction_generated (propone sustitucion/adicion/
    reestructuracion, ligada al finding y la regulacion via su
    recomendacion ya generada).
  - evidencia_insuficiente -> evidence_required (nunca se inventan hechos,
    fechas, firmas o resultados; se genera una solicitud de evidencia).
  - no_aplica -> not_applicable_justification_required (requiere
    justificacion formal aprobada por QA/Validacion, no autoaprobacion).
  - cualquier otro valor -> human_decision_required (conservador).
"""

from __future__ import annotations

import hashlib
import io
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt, RGBColor

_ESTADO_TO_CORRECTION_STATUS = {
    "cumple_parcialmente": "correction_generated",
    "no_cumple": "correction_generated",
    "evidencia_insuficiente": "evidence_required",
    "no_aplica": "not_applicable_justification_required",
}


def _finding_id(finding: dict, index: int) -> str:
    basis = f"{finding.get('agente_responsable')}|{finding.get('documento')}|{finding.get('requisito_regulatorio')}"
    return f"F-{hashlib.sha256(basis.encode('utf-8')).hexdigest()[:10]}-{index:04d}"


def classify_correction(finding: dict) -> str:
    return _ESTADO_TO_CORRECTION_STATUS.get(finding.get("estado"), "human_decision_required")


def build_finding_correction_matrix(report_data: dict) -> list[dict]:
    """Matriz completa sobre los 267 findings reales del RC canonico
    (report_data proviene de gmpai_artifact_service.build_final_report_data(),
    ya leido del RC aprobado — esta funcion no vuelve a leer disco)."""
    matrix = []
    idx = 0
    for aid, rows in report_data.get("matrices", {}).items():
        agent_meta = report_data.get("agents", {}).get(aid, {})
        for r in rows:
            idx += 1
            status = classify_correction(r)
            matrix.append({
                "finding_id": _finding_id(r, idx),
                "documento": r.get("documento"),
                "agente": aid,
                "agent_version": agent_meta.get("agent_version"),
                "clasificacion": r.get("estado"),
                "severidad": r.get("severidad"),
                "confianza": r.get("confianza"),
                "requisito_regulatorio": r.get("requisito_regulatorio"),
                "brecha": r.get("brecha"),
                "correccion_posible": status == "correction_generated",
                "motivo": {
                    "correction_generated": "estado con recomendacion concreta del agente, sustentable como propuesta trazable",
                    "evidence_required": "evidencia insuficiente — no se inventan hechos, se solicita evidencia",
                    "not_applicable_justification_required": "requiere justificacion formal QA/Validacion, no autoaprobacion",
                    "human_decision_required": "estado no clasificado automaticamente, requiere decision humana",
                }[status],
                "accion_propuesta": r.get("recomendacion") or "no_disponible",
                "estado_correccion": status,
                "revision_humana": True,
                "responsable": "QA/Validacion",
            })
    return matrix


def summarize_correction_matrix(matrix: list[dict]) -> dict:
    total = len(matrix)
    by_status: dict[str, int] = {}
    docs = set()
    for row in matrix:
        by_status[row["estado_correccion"]] = by_status.get(row["estado_correccion"], 0) + 1
        docs.add(row["documento"])
    return {
        "findings_totales": total,
        "documentos_afectados": len(docs),
        "corregibles": by_status.get("correction_generated", 0),
        "evidencia_requerida": by_status.get("evidence_required", 0),
        "no_aplica_requiere_justificacion": by_status.get("not_applicable_justification_required", 0),
        "decision_humana_requerida": by_status.get("human_decision_required", 0),
    }


def _sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def build_document_correction_draft_docx(
    documento: str, findings_for_doc: list[dict], agent_versions: dict, source_sha256: str | None = None,
) -> tuple[bytes, str]:
    """Borrador CONSOLIDADO de un unico documento: todas sus irregularidades
    en un solo DOCX (no uno por finding), separadas por estado_correccion.
    No toca el original; no marca nada como approved/compliant."""
    doc = Document()

    title = doc.add_heading(f"Propuesta de Remediacion Consolidada — DRAFT", level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    warn = doc.add_paragraph()
    warn_run = warn.add_run(
        "ESTADO: DRAFT — requiere revision humana (QA/Validacion). Ningun hallazgo se "
        "marca como corregido, aprobado o conforme automaticamente. El documento original "
        "NO fue modificado ni sobrescrito."
    )
    warn_run.bold = True
    warn_run.font.color.rgb = RGBColor(0xB0, 0x40, 0x00)
    warn_run.font.size = Pt(10)

    doc.add_heading("Documento evaluado", level=2)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for label, value in [
        ("Documento original", documento),
        ("SHA-256 original", source_sha256 or "no_disponible"),
        ("Findings incluidos", str(len(findings_for_doc))),
        ("Fecha de generacion", datetime.now(timezone.utc).isoformat()),
        ("Version del borrador", "v1 (draft)"),
    ]:
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value)

    grouped: dict[str, list[dict]] = {}
    for f in findings_for_doc:
        status = classify_correction(f)
        grouped.setdefault(status, []).append(f)

    section_titles = {
        "correction_generated": "Correcciones propuestas (cumple parcialmente / no cumple)",
        "evidence_required": "Solicitudes de evidencia (evidencia insuficiente)",
        "not_applicable_justification_required": "Requieren justificacion formal (no aplica)",
        "human_decision_required": "Requieren decision humana directa",
    }

    ct = doc.add_table(rows=1, cols=6)
    ct.style = "Light Grid Accent 1"
    hdr = ct.rows[0].cells
    for i, h in enumerate(["req.", "estado", "severidad", "brecha", "accion propuesta", "agente"]):
        hdr[i].text = h

    for status in ("correction_generated", "evidence_required",
                   "not_applicable_justification_required", "human_decision_required"):
        rows = grouped.get(status, [])
        if not rows:
            continue
        doc.add_heading(section_titles[status], level=2)
        for f in rows:
            row = ct.add_row().cells
            row[0].text = str(f.get("requisito_regulatorio") or "")
            row[1].text = str(f.get("estado") or "")
            row[2].text = str(f.get("severidad") or "")
            row[3].text = str(f.get("brecha") or "")
            row[4].text = str(f.get("recomendacion") or "no_disponible")
            row[5].text = str(f.get("agente_responsable") or "")
            p = doc.add_paragraph()
            p.add_run(f"{f.get('requisito_regulatorio')} ").bold = True
            p.add_run(f"[{f.get('estado')} / {f.get('severidad')}] — {f.get('agente_responsable')} "
                      f"(agent_version {agent_versions.get(f.get('agente_responsable'), {}).get('agent_version', 'no_disponible')})")
            doc.add_paragraph(f"Brecha: {f.get('brecha') or 'no_disponible'}")
            if status == "correction_generated":
                doc.add_paragraph(
                    "Texto propuesto (DRAFT, requiere validacion humana — NO es un hecho "
                    f"verificado): {f.get('recomendacion') or 'no_disponible'}"
                )
            elif status == "evidence_required":
                doc.add_paragraph(
                    "Solicitud de evidencia: no se genera texto sustitutivo. Se requiere "
                    "aportar evidencia real (documento, firma, fecha, resultado) — no se "
                    "asume ni se inventa su existencia."
                )
            doc.add_paragraph(f"Confianza del hallazgo: {f.get('confianza') or 'no_disponible'}")
            doc.add_paragraph("")

    doc.add_heading("Changelog", level=2)
    chg = doc.add_table(rows=1, cols=3)
    chg.style = "Light Grid Accent 1"
    hdr2 = chg.rows[0].cells
    hdr2[0].text, hdr2[1].text, hdr2[2].text = "Version", "Fecha", "Cambio"
    row = chg.add_row().cells
    row[0].text = "v1 (draft)"
    row[1].text = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    row[2].text = f"Generacion inicial consolidada a partir de {len(findings_for_doc)} findings reales del RC canonico."

    doc.add_heading("Limitaciones declaradas", level=2)
    doc.add_paragraph(
        "Este documento NO reemplaza el original ni constituye evidencia de cumplimiento. "
        "Las secciones 'evidencia insuficiente' NO incluyen texto sustitutivo — nunca se "
        "inventan fechas, firmas ni resultados. Requiere decision (accept/reject/"
        "request_changes) de QA/Validacion via el mecanismo oficial de Capa 9."
    )

    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    return data, _sha256_bytes(data)
