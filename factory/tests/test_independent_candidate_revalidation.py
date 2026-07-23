"""
Tests -- W5 V2 Fase O: factory.services.independent_candidate_revalidation
(AGT-RVL).

Cubre: un cambio incluido y realmente presente en el candidato -> CLOSED;
un cambio excluido -> OPEN; contenido original eliminado -> gap nuevo
detectado; consistencia entre redline/matriz/manifest verificada de forma
independiente (sin llamar a resolve_package_changes de Fase K).
"""
import io
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from docx import Document

from factory.services.candidate_document_generator import generate_candidate_document
from factory.services.independent_candidate_revalidation import (
    revalidate_document,
)

STRUCTURE = {
    "total_paginas": 20,
    "secciones": [
        {"numero": "1", "titulo": "Introduction", "pagina_inicio": 1, "parrafos": ["Texto original 1."]},
        {"numero": "2", "titulo": "Security", "pagina_inicio": 10, "parrafos": ["Texto original 2."]},
    ],
    "texto_previo_a_primera_seccion": ["Portada."],
    "toc_anchored": True,
}


def _change(change_id, proposed_content="Contenido nuevo propuesto.", requirement_id="21_CFR_11.10(a)"):
    return {
        "change_id": change_id, "requirement_id": requirement_id, "change_type": "CONTENT_ADDITION",
        "document_location": "seccion 2", "proposed_content": proposed_content,
        "citations": [{"page_start": 10}],
    }


def _doc_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestRevalidateChangeViaFullDocument:

    def test_included_and_present_change_is_closed(self):
        change = _change("C1")
        candidate = generate_candidate_document(STRUCTURE, [change])
        candidate_bytes = _doc_bytes(candidate)
        result = revalidate_document(
            structure=STRUCTURE, changes=[change], included_change_ids=["C1"],
            candidate_document_bytes=candidate_bytes, redline_change_ids=["C1"],
            matrix_change_ids=["C1"], manifest_artifact_hashes={"candidate": "a" * 64},
            required_manifest_artifacts=["candidate"],
        )
        assert result.gap_results[0].gap_status == "CLOSED"
        assert result.revalidation_passed is True

    def test_excluded_change_is_open(self):
        change = _change("C1")
        candidate = generate_candidate_document(STRUCTURE, [])  # nunca se incluyo
        candidate_bytes = _doc_bytes(candidate)
        result = revalidate_document(
            structure=STRUCTURE, changes=[change], included_change_ids=[],
            candidate_document_bytes=candidate_bytes, redline_change_ids=[],
            matrix_change_ids=["C1"], manifest_artifact_hashes={"candidate": "a" * 64},
            required_manifest_artifacts=["candidate"],
        )
        assert result.gap_results[0].gap_status == "OPEN"
        assert result.revalidation_passed is False

    def test_declared_included_but_actually_absent_stays_open_not_fabricated_closed(self):
        """Si insertion_manifest dice que se incluyo pero el texto real no
        esta en el candidato reabierto, NUNCA se declara CLOSED --
        detecta desincronizacion real entre lo declarado y el documento."""
        change = _change("C1", proposed_content="Este texto nunca se genero de verdad.")
        candidate = generate_candidate_document(STRUCTURE, [])  # documento sin el cambio real
        candidate_bytes = _doc_bytes(candidate)
        result = revalidate_document(
            structure=STRUCTURE, changes=[change], included_change_ids=["C1"],  # declarado incluido, mintiendo
            candidate_document_bytes=candidate_bytes, redline_change_ids=["C1"],
            matrix_change_ids=["C1"], manifest_artifact_hashes={"candidate": "a" * 64},
            required_manifest_artifacts=["candidate"],
        )
        assert result.gap_results[0].gap_status == "OPEN"


class TestNewGapDetection:

    def test_no_new_gaps_when_original_content_preserved(self):
        change = _change("C1")
        candidate = generate_candidate_document(STRUCTURE, [change])
        result = revalidate_document(
            structure=STRUCTURE, changes=[change], included_change_ids=["C1"],
            candidate_document_bytes=_doc_bytes(candidate), redline_change_ids=["C1"],
            matrix_change_ids=["C1"], manifest_artifact_hashes={"candidate": "a" * 64},
            required_manifest_artifacts=["candidate"],
        )
        assert result.new_gaps_introduced == []

    def test_detects_removed_original_content_as_new_gap(self):
        """Simula un candidato que perdio contenido original real --
        gate obligatorio del plan: 'sin eliminacion de contenido
        requerido'."""
        doc = Document()
        doc.add_heading("Documento truncado", level=1)
        doc.add_paragraph("Solo queda esto, se perdio el resto.")
        broken_bytes = _doc_bytes(doc)
        result = revalidate_document(
            structure=STRUCTURE, changes=[], included_change_ids=[],
            candidate_document_bytes=broken_bytes, redline_change_ids=[],
            matrix_change_ids=[], manifest_artifact_hashes={"candidate": "a" * 64},
            required_manifest_artifacts=["candidate"],
        )
        assert len(result.new_gaps_introduced) == 2  # los 2 parrafos originales de STRUCTURE
        assert result.revalidation_passed is False


class TestArtifactConsistency:

    def test_mismatched_redline_and_matrix_fails_consistency(self):
        change = _change("C1")
        candidate = generate_candidate_document(STRUCTURE, [change])
        result = revalidate_document(
            structure=STRUCTURE, changes=[change], included_change_ids=["C1"],
            candidate_document_bytes=_doc_bytes(candidate),
            redline_change_ids=["C1", "C2_FANTASMA"],  # inconsistente
            matrix_change_ids=["C1"], manifest_artifact_hashes={"candidate": "a" * 64},
            required_manifest_artifacts=["candidate"],
        )
        assert result.artifacts_consistent is False
        assert result.revalidation_passed is False

    def test_missing_manifest_artifact_fails_consistency(self):
        change = _change("C1")
        candidate = generate_candidate_document(STRUCTURE, [change])
        result = revalidate_document(
            structure=STRUCTURE, changes=[change], included_change_ids=["C1"],
            candidate_document_bytes=_doc_bytes(candidate), redline_change_ids=["C1"],
            matrix_change_ids=["C1"], manifest_artifact_hashes={},
            required_manifest_artifacts=["candidate", "redline"],
        )
        assert result.artifacts_consistent is False

    def test_invalid_hash_format_fails_hashes_valid(self):
        change = _change("C1")
        candidate = generate_candidate_document(STRUCTURE, [change])
        result = revalidate_document(
            structure=STRUCTURE, changes=[change], included_change_ids=["C1"],
            candidate_document_bytes=_doc_bytes(candidate), redline_change_ids=["C1"],
            matrix_change_ids=["C1"], manifest_artifact_hashes={"candidate": "not-a-real-hash"},
            required_manifest_artifacts=["candidate"],
        )
        assert result.all_hashes_valid is False


class TestDocumentDoesNotOpen:

    def test_corrupt_bytes_marks_document_as_not_opening(self):
        result = revalidate_document(
            structure=STRUCTURE, changes=[], included_change_ids=[],
            candidate_document_bytes=b"esto no es un docx valido",
            redline_change_ids=[], matrix_change_ids=[], manifest_artifact_hashes={},
            required_manifest_artifacts=[],
        )
        assert result.document_opens_correctly is False
        assert result.revalidation_passed is False


class TestNeverAutoAssignsImplementationVerification:

    def test_no_gap_result_is_ever_implementation_verification_required(self):
        """Regla dura del plan: IMPLEMENTATION_VERIFICATION_REQUIRED
        nunca se asigna automaticamente -- requiere verificar el sistema
        fisico, decision siempre humana."""
        changes = [_change("C1"), _change("C2")]
        candidate = generate_candidate_document(STRUCTURE, changes)
        result = revalidate_document(
            structure=STRUCTURE, changes=changes, included_change_ids=["C1", "C2"],
            candidate_document_bytes=_doc_bytes(candidate), redline_change_ids=["C1", "C2"],
            matrix_change_ids=["C1", "C2"], manifest_artifact_hashes={"candidate": "a" * 64},
            required_manifest_artifacts=["candidate"],
        )
        assert all(r.gap_status != "IMPLEMENTATION_VERIFICATION_REQUIRED" for r in result.gap_results)
