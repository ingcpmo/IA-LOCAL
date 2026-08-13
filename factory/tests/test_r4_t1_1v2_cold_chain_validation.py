"""R4-T1.1v2 §2 (docs_plan/R4_T1_1v2_DESBLOQUEO_Y_VALIDACION_FRIA.md) --
validacion en frio de la cadena completa sobre RW-0005 (piloto confirmado
por Cesar el 2026-08-13, ver §1 del mismo plan). CERO llamadas LLM: el
hallazgo de entrada es SINTETICO (marcado DRY_RUN_VALIDATION, cola
aislada via el fixture autouse isolated_review_queue, nunca puede
colisionar con el PROVISIONAL_GAP real de 21_CFR_11.10(d) ya pendiente en
factory/layer9/review_queue.jsonl -- ese archivo real ni se abre en este
test). La directiva es de autoria humana real de Cesar (2026-08-13,
sesion interactiva) -- este test NUNCA genera proposed_text/rationale,
solo los transporta literalmente."""
from __future__ import annotations

import hashlib
from pathlib import Path

import pytest

from factory.core import audit_writer
from factory.layer9 import human_review_queue as hrq
from factory.services import paths
from factory.services import remediation_directive as rd
from factory.services import remediation_directive_dispatch as dispatch
from factory.services import remediation_package_service as pkg_svc
from factory.services.governed_candidate_document_pipeline import generate_governed_candidate_from_pdf
from factory.services.candidate_document_generator import NO_APROBADO_MARK, verify_document_conformance
from factory.regulatory.corpus_runner import _resolve_document_path

RW_0005_PATH = Path(
    "/home/ing_cpmo/GMPAI/source/Rockwell/215115305 SCADA-PCS Misc PLC System FS_v1.2.pdf"
)

pytestmark = pytest.mark.skipif(
    not RW_0005_PATH.exists(), reason="requiere el PDF real de Rockwell, no presente en este entorno")


@pytest.fixture(autouse=True)
def _isolated_directives(tmp_path, monkeypatch):
    directives_file = tmp_path / "remediation_directives_test.jsonl"
    monkeypatch.setattr(rd, "DIRECTIVES_FILE", directives_file)
    yield directives_file


@pytest.fixture(autouse=True)
def _isolated_packages_and_audit(tmp_path, monkeypatch):
    monkeypatch.setattr(paths, "REMEDIATION_PACKAGES_BASE", tmp_path / "remediation_packages")
    monkeypatch.setattr(audit_writer, "AUDIT_FILE", tmp_path / "audit" / "test_factory_audit.jsonl")
    monkeypatch.setattr(audit_writer, "_last_entry_hash", None)
    yield


DRY_RUN_RC_ID = "finding-DRY_RUN_VALIDATION-RW-0005-21_CFR_11.10(d)"


def _make_synthetic_confirmed_finding() -> dict:
    """Hallazgo SINTETICO bien formado (misma forma que finding_review_v2
    real) para 21_CFR_11.10(d) sobre RW-0005 -- marcado DRY_RUN_VALIDATION,
    escrito por este test (aislado por isolated_review_queue) directamente
    con status=confirmed. Ubicacion real (Seccion 4 Security, pag 40-44) y
    cita real del catalogo -- lo unico sintetico es que este hallazgo en
    particular no vino de una corrida LLM real."""
    return {
        "schema_version": "finding_review_v2",
        "rc_id": DRY_RUN_RC_ID,
        "entry_type": "finding_review",
        "project_id": "RW-0005",
        "enqueued_at": "2026-08-13T00:00:00+00:00",
        "status": "confirmed",
        "summary": {
            "run_id": "DRY_RUN_VALIDATION-R4-T1.1v2",
            "requirement_id": "21_CFR_11.10(d)",
            "document_id": "RW-0005",
            "page": 40,
            "evidence_quote": "",
            "conclusion": "PROVISIONAL_GAP",
            "review_flags": ["DRY_RUN_VALIDATION"],
            "agent_id": "DRY_RUN_VALIDATION",
            "candidates": [],
        },
        "reviewer": "cesar",
        "reviewed_at": "2026-08-13T00:00:00+00:00",
        "human_confirmed_evidence": {
            "page": None, "quote": "DRY_RUN_VALIDATION -- R4-T1.1v2 SS2",
            "confirmed_by": "cesar", "confirmed_at": "2026-08-13T00:00:00+00:00",
        },
    }


def _write_synthetic_finding_to_isolated_queue():
    hrq.REVIEW_QUEUE_FILE.parent.mkdir(parents=True, exist_ok=True)
    with open(hrq.REVIEW_QUEUE_FILE, "a", encoding="utf-8") as f:
        import json
        f.write(json.dumps(_make_synthetic_confirmed_finding(), ensure_ascii=False) + "\n")


def test_synthetic_finding_isolated_never_touches_real_queue(tmp_path):
    real_queue = Path(__file__).resolve().parents[1] / "layer9" / "review_queue.jsonl"
    real_bytes_before = real_queue.read_bytes()

    _write_synthetic_finding_to_isolated_queue()

    assert hrq.REVIEW_QUEUE_FILE != real_queue
    assert hrq.REVIEW_QUEUE_FILE.parent == tmp_path
    entry = hrq.get_entry(DRY_RUN_RC_ID)
    assert entry is not None
    assert entry["summary"]["review_flags"] == ["DRY_RUN_VALIDATION"]

    # el hallazgo real PROVISIONAL_GAP de 21_CFR_11.10(d) sigue intacto,
    # el archivo real ni siquiera se toco durante este test
    assert real_queue.read_bytes() == real_bytes_before


CESAR_DIRECTIVE = dict(
    finding_rc_id=DRY_RUN_RC_ID,
    change_type="ADD",
    proposed_text=(
        "El sistema debe registrar y revisar periodicamente la lista de "
        "individuos autorizados y sus niveles de acceso en FactoryTalk "
        "Security, con revocacion documentada al cambiar roles o cesar "
        "funciones."
    ),
    target_location={"page_start": 40, "page_end": 44, "section": "4 Security"},
    regulatory_citation=["21_CFR_11.10(d)"],
    rationale=(
        "El documento describe niveles de acceso pero no exige la revision "
        "periodica ni la revocacion documentada que 21 CFR 11.10(d) implica."
    ),
    authored_by_id="cesar",
    authored_by_display_name="Cesar",
)


def test_directive_authored_by_cesar_is_accepted_and_dispatches_to_high_risk_change():
    _write_synthetic_finding_to_isolated_queue()

    directive = rd.propose_remediation_directive(**CESAR_DIRECTIVE)
    assert directive["status"] == "SUBMITTED"
    assert directive["authored_by_id"] == "cesar"
    assert directive["document_id"] == "RW-0005"
    assert directive["document_sha256"] == hashlib.sha256(RW_0005_PATH.read_bytes()).hexdigest()

    mapped = dispatch.dispatch_directive_to_remediation(directive, run_id="DRY_RUN_VALIDATION-R4-T1.1v2")
    change = mapped.change
    print("\n--- RemediationChange real (Ruta D) ---")
    for k, v in change.items():
        print(f"{k}: {v}")
    print("--- risk_factors ---", mapped.risk_factors)
    print("--- confidence_factors ---", mapped.confidence_factors)


# ── §2.2 cadena completa + §2.3 criterios a-h ───────────────────────────────

import hashlib as _hashlib
import subprocess
from datetime import datetime, timezone

import pytest as _pytest

DRY_RUN_OUTPUT_DIR = (
    Path(__file__).resolve().parents[1] / "regulatory" / "pilot_run"
    / "dry_run_validation_r4_t1_1v2"
)
PROJECT_ID = "RW-0005-DRY-RUN-VALIDATION"
PACKAGE_ID = "PKG-DRY-RUN-R4-T1-1v2"


def _sha256_bytes(data: bytes) -> str:
    return _hashlib.sha256(data).hexdigest()


def _artifact_ref(kind: str, path: Path) -> dict:
    classification = {
        "candidate_document": "CANDIDATE_DRAFT", "redline_document": "REDLINE",
    }[kind]
    data = path.read_bytes()
    return {
        "artifact_id": f"ART-{kind}-{PACKAGE_ID}", "storage_location": str(path),
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "sha256": _sha256_bytes(data), "size_bytes": len(data),
        "classification": classification,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }


def _basis() -> dict:
    return {
        "requirements_applicable": 1,
        "coverage_complete_by_requirement": {"21_CFR_11.10(d)": True},
        "expected_chunks": 1, "evaluated_chunks": 1,
        "execution_errors": 0, "rejected_records": 0,
    }


def test_full_cold_chain_rw0005_directive_to_traceable_candidate(tmp_path):
    """§2.2 completo sobre RW-0005: directiva (Acto 2, Cesar) ->
    remediation_directive_dispatch -> Ruta D -> RemediationChange ->
    remediation_package_service.create_package -> generador de candidato
    -> redline -> manifest -> trazabilidad. Evalua los 8 criterios de
    aceptacion de §2.3 (a-h). CERO llamadas LLM."""
    commit_sha = subprocess.run(
        ["git", "rev-parse", "HEAD"], cwd=Path(__file__).resolve().parents[2],
        capture_output=True, text=True, check=True,
    ).stdout.strip()

    pdf_path, doc_sha_before = _resolve_document_path("RW-0005")
    sha_before = _sha256_bytes(pdf_path.read_bytes())

    # 1) hallazgo sintetico aislado, confirmado
    _write_synthetic_finding_to_isolated_queue()

    # 2) directiva real de autoria humana (Cesar, Acto 2)
    directive = rd.propose_remediation_directive(**CESAR_DIRECTIVE)

    # 3) Ruta D: directiva -> RemediationChange real
    mapped = dispatch.dispatch_directive_to_remediation(directive, run_id="DRY_RUN_VALIDATION-R4-T1.1v2")
    change = mapped.change
    change_id = change["change_id"]
    assert change["change_risk"] == "HIGH_RISK"

    # 4) v1: candidato/redline generados con el HIGH_RISK todavia sin
    #    excepcion revisada -- prueba real de que NUNCA se incorpora en
    #    silencio (regla dura del plan).
    package_state_v1 = {
        "changes": {change_id: change}, "exceptions": {},
        "medium_risk_batch_decisions": {}, "package_decision": None,
    }
    result_v1 = generate_governed_candidate_from_pdf(pdf_path, package_state_v1)
    assert result_v1.included_change_ids == []
    assert result_v1.excluded_change_ids == [change_id]
    assert "EXCEPTION_REQUIRED" in result_v1.exclusion_reasons[change_id]

    DRY_RUN_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    candidate_v1_path = DRY_RUN_OUTPUT_DIR / "v1_candidate_EXCLUDED_pending_exception.docx"
    redline_v1_path = DRY_RUN_OUTPUT_DIR / "v1_redline_EXCLUDED_pending_exception.docx"
    result_v1.candidate_document.save(str(candidate_v1_path))
    result_v1.redline_document.save(str(redline_v1_path))

    artifacts_v1 = {
        "candidate_document": _artifact_ref("candidate_document", candidate_v1_path),
        "redline_document": _artifact_ref("redline_document", redline_v1_path),
    }
    pkg_v1 = pkg_svc.create_package(
        project_id=PROJECT_ID, package_id=PACKAGE_ID, package_version=1,
        changes=[change], artifacts=artifacts_v1, automatic_evaluation_basis=_basis(),
        generation_commit_sha=commit_sha,
    )
    assert pkg_v1["status"] == "AWAITING_HUMAN_EXCEPTION_REVIEW"
    assert pkg_v1["changes"]["high_risk"] == [change_id]

    # 5) excepcion HIGH_RISK -- decision MECANICA de Capa 8 para validar el
    #    pipeline en un paquete DRY_RUN aislado que nunca se libera ni se
    #    usa como base de un candidato real. NO es una aceptacion de riesgo
    #    GxP genuina -- esa autoridad sigue siendo de QA/Cesar/Capa 9
    #    (CLAUDE.md). Queda registrado explicitamente en la justificacion
    #    para que nadie confunda esta entrada con un sign-off real.
    exc = pkg_svc.record_exception_review(
        project_id=PROJECT_ID, package_id=PACKAGE_ID, package_version=1,
        change_id=change_id, human_review_decision="accept_dry_run_pipeline_validation",
        responsible="capa8_claude_code",
        justification=(
            "DRY_RUN_VALIDATION R4-T1.1v2 SS2 -- decision mecanica de Capa 8 "
            "para ejercitar la cadena de excepcion HIGH_RISK de un paquete "
            "aislado que nunca se libera ni se usa como candidato real. NO "
            "constituye una aceptacion de riesgo GxP genuina -- esa autoridad "
            "sigue siendo de QA/Cesar/Capa 9 (CLAUDE.md)."
        ),
    )
    assert exc["status"] == "REVIEWED"
    state_v1_after = pkg_svc._read_state(PROJECT_ID, PACKAGE_ID, 1)
    assert state_v1_after["package"]["status"] == "AWAITING_PACKAGE_DECISION"
    assert state_v1_after["package"]["human_exception_review_complete"] is True

    # 6) v2: mismo change_id, package_state ahora refleja la excepcion
    #    aceptada -> AUTO_APPLIED_TO_DRAFT -> el candidato SI incluye el
    #    cambio de Cesar.
    result_v2 = generate_governed_candidate_from_pdf(pdf_path, state_v1_after)
    assert result_v2.included_change_ids == [change_id]
    assert result_v2.excluded_change_ids == []

    candidate_v2_path = DRY_RUN_OUTPUT_DIR / "v2_candidate_INCLUDED.docx"
    redline_v2_path = DRY_RUN_OUTPUT_DIR / "v2_redline_INCLUDED.docx"
    result_v2.candidate_document.save(str(candidate_v2_path))
    result_v2.redline_document.save(str(redline_v2_path))

    artifacts_v2 = {
        "candidate_document": _artifact_ref("candidate_document", candidate_v2_path),
        "redline_document": _artifact_ref("redline_document", redline_v2_path),
    }
    pkg_v2 = pkg_svc.create_package(
        project_id=PROJECT_ID, package_id=PACKAGE_ID, package_version=2,
        changes=[change], artifacts=artifacts_v2, automatic_evaluation_basis=_basis(),
        generation_commit_sha=commit_sha,
    )
    exc_v2 = pkg_svc.record_exception_review(
        project_id=PROJECT_ID, package_id=PACKAGE_ID, package_version=2,
        change_id=change_id, human_review_decision="accept_dry_run_pipeline_validation",
        responsible="capa8_claude_code",
        justification="Misma decision mecanica que v1 (ver justificacion de v1) -- version regenerada tras excepcion aceptada, DRY_RUN_VALIDATION, sin efecto GxP real.",
    )
    state_v2_after = pkg_svc._read_state(PROJECT_ID, PACKAGE_ID, 2)

    # ── criterio (a): el candidato se genera y ABRE desde disco ─────────
    from docx import Document as _Document
    reopened_candidate = _Document(str(candidate_v2_path))
    reopened_texts = [p.text for p in reopened_candidate.paragraphs]
    crit_a = candidate_v2_path.exists() and redline_v2_path.exists() and len(reopened_texts) > 0
    print(f"\ncriterio (a) candidato se genera y abre desde disco: {'PASA' if crit_a else 'FALLA'}")
    assert crit_a

    # ── criterio (b): marca BORRADOR -- NO APROBADO visible ─────────────
    crit_b = NO_APROBADO_MARK in reopened_texts
    reopened_redline = _Document(str(redline_v2_path))
    crit_b = crit_b and NO_APROBADO_MARK in [p.text for p in reopened_redline.paragraphs]
    print(f"criterio (b) marca NO-APROBADO presente en candidato y redline: {'PASA' if crit_b else 'FALLA'}")
    assert crit_b

    # ── criterio (c): el redline refleja EXACTAMENTE el cambio (reapertura) ─
    conformance = verify_document_conformance(str(redline_v2_path), [change], result_v2.insertion_manifest)
    crit_c = len(conformance) == 1 and conformance[0]["status"] == "DOCUMENT_CONFORMANCE"
    print(f"criterio (c) redline == cambio de la directiva (DOCUMENT_CONFORMANCE): "
          f"{'PASA' if crit_c else 'FALLA'} -- {conformance}")
    assert crit_c
    proposed_text_literal_present = any(CESAR_DIRECTIVE["proposed_text"] in t for t in reopened_texts)
    print(f"  texto literal de Cesar presente en el candidato: {proposed_text_literal_present}")
    assert proposed_text_literal_present

    # ── criterio (d): documento ORIGINAL intacto (SHA-256 antes/despues) ─
    sha_after = _sha256_bytes(pdf_path.read_bytes())
    crit_d = sha_before == sha_after == doc_sha_before
    print(f"criterio (d) original intacto (SHA-256): {'PASA' if crit_d else 'FALLA'} "
          f"antes={sha_before[:12]}.. despues={sha_after[:12]}..")
    assert crit_d

    # ── criterio (e): trazabilidad completa cambio -> directiva -> hallazgo ─
    manifest_entry = result_v2.insertion_manifest[0]
    crit_e = (
        manifest_entry["change_id"] == change_id == change["finding_id"] == directive["directive_id"]
        and directive["finding_rc_id"] == DRY_RUN_RC_ID
        and hrq.get_entry(DRY_RUN_RC_ID) is not None
    )
    print(f"criterio (e) trazabilidad cambio->directiva->hallazgo: {'PASA' if crit_e else 'FALLA'}")
    print(f"  change_id={change_id}")
    print(f"  directive_id={directive['directive_id']}")
    print(f"  finding_rc_id={directive['finding_rc_id']}")
    assert crit_e

    # ── criterio (f): manifest lista todos los artefactos con sus hashes ─
    crit_f = (
        all("sha256" in a and a["sha256"] for a in artifacts_v2.values())
        and len(result_v2.insertion_manifest) == 1
        and result_v2.insertion_manifest[0]["proposed_content_sha256"] == _hashlib.sha256(
            change["proposed_content"].encode("utf-8")).hexdigest()
    )
    print(f"criterio (f) manifest de artefactos con hashes: {'PASA' if crit_f else 'FALLA'}")
    for kind, art in artifacts_v2.items():
        print(f"  {kind}: sha256={art['sha256'][:16]}.. size={art['size_bytes']}B")
    assert crit_f

    # ── criterio (g): NUNCA se alcanza create_release_record() ──────────
    # create_release_record() existe en remediation_package_service.py
    # (funcion real del servicio) pero NO esta conectada a ningun endpoint
    # HTTP (factory/api/routes/remediation_packages.py lo declara a
    # proposito: "DOCUMENT_RELEASED permanece false y PRODUCTION_ENABLEMENT
    # permanece BLOCKED"). Este test, ademas, nunca la invoca.
    routes_src = (
        Path(__file__).resolve().parents[1] / "api" / "routes" / "remediation_packages.py"
    ).read_text(encoding="utf-8")
    crit_g = "create_release_record" not in routes_src.replace(
        "create_release_record() existe en el servicio pero no se conecta", ""
    )
    print(f"criterio (g) create_release_record() sin endpoint conectado: {'PASA' if crit_g else 'FALLA'}")
    assert crit_g
    assert state_v2_after["package_decision"] is None
    assert "releases" not in state_v2_after
    assert not any(f.name.startswith("releases") for f in DRY_RUN_OUTPUT_DIR.glob("*"))

    print(f"\npkg_v1 status final (tras record_exception_review): {state_v1_after['package']['status']}")
    print(f"pkg_v2 status final (tras record_exception_review): {state_v2_after['package']['status']}, "
          f"human_exception_review_complete={state_v2_after['package']['human_exception_review_complete']}")
    print(f"artefactos guardados en: {DRY_RUN_OUTPUT_DIR}")


def test_delete_change_type_still_rejected():
    """criterio (h), primera mitad: change_type=DELETE sigue rechazado."""
    _write_synthetic_finding_to_isolated_queue()
    directive = dict(CESAR_DIRECTIVE)
    directive["change_type"] = "DELETE"
    directive["original_text"] = "texto que no importa para esta prueba"
    with _pytest.raises(rd.RemediationDirectiveError):
        rd.propose_remediation_directive(**directive)


def test_directive_without_citation_still_rejected():
    """criterio (h), segunda mitad (sin cita)."""
    _write_synthetic_finding_to_isolated_queue()
    directive = dict(CESAR_DIRECTIVE)
    directive["regulatory_citation"] = []
    with _pytest.raises(rd.RemediationDirectiveError):
        rd.propose_remediation_directive(**directive)


def test_directive_with_non_anchoring_original_text_still_rejected():
    """criterio (h), tercera parte (original_text que no ancla)."""
    _write_synthetic_finding_to_isolated_queue()
    directive = dict(CESAR_DIRECTIVE)
    directive["change_type"] = "REPLACE"
    directive["original_text"] = "esta frase no existe literalmente en las paginas 40-44 del PDF real"
    with _pytest.raises(rd.RemediationDirectiveError):
        rd.propose_remediation_directive(**directive)


def test_directive_with_stale_document_sha_still_rejected():
    """criterio (h), cuarta parte (sha desactualizado) -- se ejercita en
    dispatch, que re-verifica el SHA-256 vivo contra el registrado en la
    directiva (ver remediation_directive_dispatch.dispatch_directive_to_remediation)."""
    _write_synthetic_finding_to_isolated_queue()
    directive = rd.propose_remediation_directive(**CESAR_DIRECTIVE)
    stale = dict(directive)
    stale["document_sha256"] = "0" * 64
    with _pytest.raises(dispatch.DirectiveNotDispatchable):
        dispatch.dispatch_directive_to_remediation(stale, run_id="DRY_RUN_VALIDATION-R4-T1.1v2")
