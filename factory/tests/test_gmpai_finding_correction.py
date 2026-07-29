"""
Tests — clasificacion finding -> correccion y borrador consolidado por
documento, sobre datos REALES del RC canonico de gmpai_document_validation
(no reprocesa documentos, no invoca agentes, no inventa contenido).
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from factory.services import gmpai_artifact_service as svc
from factory.services import gmpai_finding_correction_service as fcs


def _real_report_data():
    return svc.build_final_report_data()


def test_classify_correction_mapping():
    assert fcs.classify_correction({"estado": "cumple_parcialmente"}) == "correction_generated"
    assert fcs.classify_correction({"estado": "no_cumple"}) == "correction_generated"
    assert fcs.classify_correction({"estado": "evidencia_insuficiente"}) == "evidence_required"
    assert fcs.classify_correction({"estado": "no_aplica"}) == "not_applicable_justification_required"
    assert fcs.classify_correction({"estado": "algo_no_previsto"}) == "human_decision_required"
    assert fcs.classify_correction({"estado": "cumple"}) == "human_decision_required"  # no hay 0 "cumple" reales


def test_build_matrix_covers_every_real_finding_without_losing_any():
    """La matriz cubre TODOS los findings reales del reporte, sean los que
    sean. El `== 267` que habia aqui era redundante: el invariante util ya
    estaba escrito al lado (`len(matrix) == data["findings_total"]`), y el
    numero solo anadia una forma de romperse cuando el corpus cambie."""
    data = _real_report_data()
    matrix = fcs.build_finding_correction_matrix(data)
    assert len(matrix) == data["findings_total"]
    assert all(row["finding_id"] for row in matrix)
    assert len({row["finding_id"] for row in matrix}) == len(matrix), "finding_id debe ser unico"


def test_matrix_never_marks_evidencia_insuficiente_as_correction_generated():
    data = _real_report_data()
    matrix = fcs.build_finding_correction_matrix(data)
    for row in matrix:
        if row["clasificacion"] == "evidencia_insuficiente":
            assert row["estado_correccion"] == "evidence_required"
            assert row["correccion_posible"] is False


def test_summary_is_a_faithful_partition_of_the_real_matrix():
    """Lo que el resumen debe garantizar es que NO pierde ni duplica
    findings al agregarlos: las 4 categorias son una particion exacta de la
    matriz real, y cada una cuenta lo que dice contar.

    Antes congelaba la distribucion de un dia (267/83/184). Ese conteo se
    rompe con cualquier reanalisis legitimo del corpus, y aun asi no habria
    detectado el fallo que de verdad importa aqui: un finding contado en dos
    categorias a la vez, o perdido entre ellas."""
    data = _real_report_data()
    matrix = fcs.build_finding_correction_matrix(data)
    summary = fcs.summarize_correction_matrix(matrix)

    esperado = {
        "corregibles": "correction_generated",
        "evidencia_requerida": "evidence_required",
        "no_aplica_requiere_justificacion": "not_applicable_justification_required",
        "decision_humana_requerida": "human_decision_required",
    }
    for clave, estado in esperado.items():
        assert summary[clave] == sum(1 for r in matrix if r["estado_correccion"] == estado), clave

    assert summary["findings_totales"] == len(matrix)
    assert sum(summary[c] for c in esperado) == len(matrix), (
        "las 4 categorias deben particionar la matriz: ni un finding fuera, "
        "ni uno contado dos veces"
    )
    assert summary["documentos_afectados"] == len({r["documento"] for r in matrix})


def test_build_final_report_data_embeds_correction_summary():
    data = _real_report_data()
    assert data["finding_correction_summary"]["findings_totales"] == data["findings_total"]


def test_document_correction_draft_docx_is_valid_and_grouped_by_status():
    findings = [
        {"estado": "no_cumple", "requisito_regulatorio": "REQ-1", "severidad": "mayor",
         "brecha": "falta X", "recomendacion": "Agregar X", "agente_responsable": "fda_part11_agent",
         "confianza": "media"},
        {"estado": "evidencia_insuficiente", "requisito_regulatorio": "REQ-2", "severidad": "no_determinada",
         "brecha": "sin texto", "recomendacion": None, "agente_responsable": "alcoa_plus_agent",
         "confianza": "baja"},
        {"estado": "cumple_parcialmente", "requisito_regulatorio": "REQ-3", "severidad": "menor",
         "brecha": "falta anexo", "recomendacion": "Completar anexo", "agente_responsable": "eu_annex11_agent",
         "confianza": "media"},
    ]
    agent_versions = {
        "fda_part11_agent": {"agent_version": "1.0.0"},
        "alcoa_plus_agent": {"agent_version": "1.0.0"},
        "eu_annex11_agent": {"agent_version": "1.0.0"},
    }
    docx_bytes, sha = fcs.build_document_correction_draft_docx(
        "documento_prueba.pdf", findings, agent_versions, source_sha256="abc123")
    assert docx_bytes[:2] == b"PK"
    assert len(sha) == 64

    from docx import Document
    import io
    doc = Document(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    full_text += "\n" + "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "REQ-1" in full_text
    assert "documento_prueba.pdf" in full_text
    assert "abc123" in full_text
    # Nunca debe fabricar texto sustitutivo para evidencia insuficiente.
    assert "no se inventan fechas" in full_text or "Solicitud de evidencia" in full_text


def test_document_correction_draft_never_marks_approved_or_compliant():
    findings = [{"estado": "no_cumple", "requisito_regulatorio": "REQ-1", "severidad": "mayor",
                 "brecha": "x", "recomendacion": "y", "agente_responsable": "fda_part11_agent",
                 "confianza": "media"}]
    docx_bytes, _ = fcs.build_document_correction_draft_docx("doc.pdf", findings, {}, "sha")
    from docx import Document
    import io
    doc = Document(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs).lower()
    assert "approved" not in full_text
    assert "marcado como conforme" not in full_text
    assert "draft" in full_text
